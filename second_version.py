import os
from PyQt5.QtWidgets import QApplication, QMainWindow, QDialog, QVBoxLayout, QLabel, QLineEdit, QPushButton, QComboBox, QInputDialog, QFileDialog
from PyQt5.uic import loadUi
from docx import Document



#ОСНОВНОЕ ОКНО
class MainWindow(QMainWindow):
    def __init__(self):
        super(MainWindow, self).__init__()
        loadUi('design.ui', self)

        # чекбоксы:
        self.checkBox.setChecked(False)
        self.checkbox_two.setChecked(False)

        # анализ и данные кнопки
        self.pushButton.clicked.connect(self.on_pushButton_clicked)
        self.pushButton_2.clicked.connect(self.on_pushButton_2_clicked)

        self.checkBox.stateChanged.connect(self.on_checkBox_stateChanged)
        self.checkbox_two.stateChanged.connect(self.on_checkbox_two_stateChanged)

    def on_pushButton_clicked(self):  # кнопка анализ
        data_entry_window = DataEntryWindow()
        data_entry_window.exec_()

    def on_pushButton_2_clicked(self):  # кнопка данные
        check_data_window = CheckYourDataInFile(None)  
        check_data_window.exec_()

    def on_checkBox_stateChanged(self, state):
        if state == 2:  # state == 2, если чекбокс отмечен
            about_app_dialog = AboutApp(self)
            about_app_dialog.exec_()

    def on_checkbox_two_stateChanged(self, state):
        if state == 2:  # state == 2, если чекбокс отмечен
            about_instruction_dialog = AboutInstruction(self)
            about_instruction_dialog.exec_()


#ОКНО АНАЛИЗА
class DataEntryWindow(QDialog):
    def __init__(self):
        super().__init__()

        self.setWindowTitle("Расчет для предпринимателя")
        self.setGeometry(100, 100, 400, 300)

        layout = QVBoxLayout()

        self.label_income = QLabel("Введите данные о прибыле:")
        self.text_income = QLineEdit(self)

        self.label_expense = QLabel("Введите данные о расходах:")
        self.text_expense = QLineEdit(self)

        self.label_loan_amount = QLabel("Сумма кредита:")
        self.text_loan_amount = QLineEdit(self)

        self.label_loan_duration = QLabel("Количество месяцев обязательства:")
        self.text_loan_duration = QLineEdit(self)

        self.label_interest_rate = QLabel("Процентная ставка в месяц:")
        self.text_interest_rate = QLineEdit(self)

        self.label_net_income = QLabel("Чистая прибыль:")
        self.text_net_income = QLabel(self)

        self.label_total_payment = QLabel("Итоговая сумма по кредиту:")
        self.text_total_payment = QLabel(self)


        self.button_calculate = QPushButton("Рассчитать", self)
        self.button_calculate.clicked.connect(self.calculate_net_income_and_loan)

        self.button_clear = QPushButton("Очистить", self)
        self.button_clear.clicked.connect(self.clear_fields)

        layout.addWidget(self.label_income)
        layout.addWidget(self.text_income)
        layout.addWidget(self.label_expense)
        layout.addWidget(self.text_expense)
        layout.addWidget(self.label_loan_amount)
        layout.addWidget(self.text_loan_amount)
        layout.addWidget(self.label_loan_duration)
        layout.addWidget(self.text_loan_duration)
        layout.addWidget(self.label_interest_rate)
        layout.addWidget(self.text_interest_rate)
        layout.addWidget(self.button_calculate)
        layout.addWidget(self.label_net_income)
        layout.addWidget(self.text_net_income)
        layout.addWidget(self.label_total_payment)
        layout.addWidget(self.text_total_payment)
        layout.addWidget(self.button_clear)

        self.setLayout(layout)

    def calculate_net_income_and_loan(self):
        income = float(self.text_income.text()) if self.text_income.text() else 0.0
        expense = float(self.text_expense.text()) if self.text_expense.text() else 0.0
        net_income = income - expense
        self.text_net_income.setText(str(net_income))

        loan_amount = float(self.text_loan_amount.text()) if self.text_loan_amount.text() else 0.0
        loan_duration = int(self.text_loan_duration.text()) if self.text_loan_duration.text() else 0
        interest_rate = float(self.text_interest_rate.text()) if self.text_interest_rate.text() else 0.0

        total_payment = loan_amount * (1 + interest_rate / 100) ** loan_duration
        self.text_total_payment.setText(str(total_payment))

        total_payment = loan_amount * (1 + interest_rate / 100) ** loan_duration
        self.text_total_payment.setText(str(total_payment))

        
        result_str = f"Income: {income}\nExpense: {expense}\nNet Income: {net_income}\nLoan Amount: {loan_amount}\nLoan Duration: {loan_duration}\nInterest Rate: {interest_rate}\nTotal Payment: {total_payment}\n"

        self.save_to_file(result_str)

    def save_to_file(self, result_str):
        try:
            
            file_path = "text_results.txt"

            
            if not os.path.exists(file_path):
                with open(file_path, "w"):
                    pass

            
            with open(file_path, "a") as file:
                file.write(result_str)

        except Exception as e:
            print(f"Error saving to file: {e}")

    def clear_fields(self):
        self.text_income.clear()
        self.text_expense.clear()
        self.text_loan_amount.clear()
        self.text_loan_duration.clear()
        self.text_interest_rate.clear()
        self.text_net_income.clear()
        self.text_total_payment.clear()


#РАСЧЕТ В EXCEL (добавить запись)
class CheckYourDataInFile(QDialog):
    def __init__(self, selected_paths):
        super().__init__()

        self.setWindowTitle("Работа с файлами")
        self.setGeometry(100, 100, 400, 300)

        self.selected_paths = selected_paths

        layout = QVBoxLayout()

        self.operation_label = QLabel("Выберите операцию:")
        layout.addWidget(self.operation_label)


        self.save_result_button = QPushButton("Сохранить результат в файл", self)
        self.save_result_button.clicked.connect(self.save_result_to_file)
        layout.addWidget(self.save_result_button)

        self.operation_combo = QComboBox(self)
        self.operation_combo.addItem("Сумма")
        self.operation_combo.addItem("Разность")
        self.operation_combo.addItem("Произведение")
        self.operation_combo.addItem("Деление")
        layout.addWidget(self.operation_combo)

        self.number1_label = QLabel("Число 1:")
        self.number1_input = QLineEdit(self)
        layout.addWidget(self.number1_label)
        layout.addWidget(self.number1_input)

        self.number2_label = QLabel("Число 2:")
        self.number2_input = QLineEdit(self)
        layout.addWidget(self.number2_label)
        layout.addWidget(self.number2_input)

        self.result_label = QLabel("Результат:")
        layout.addWidget(self.result_label)

        self.result_text = QLabel(self)
        layout.addWidget(self.result_text)

        self.calculate_button = QPushButton("Выполнить операцию", self)
        self.calculate_button.clicked.connect(self.calculate_operation)
        layout.addWidget(self.calculate_button)

        self.clear_button = QPushButton("Очистить", self)
        self.clear_button.clicked.connect(self.clear_data)
        layout.addWidget(self.clear_button)

        self.setLayout(layout)
        #self.result_value = None

    def calculate_operation(self):
        operation = self.operation_combo.currentText()
        number1 = self.number1_input.text()
        number2 = self.number2_input.text()

        try:
            number1 = float(number1)
            number2 = float(number2)

            if operation == "Сумма":
                result = number1 + number2
            elif operation == "Разность":
                result = number1 - number2
            elif operation == "Произведение":
                result = number1 * number2
            elif operation == "Деление":
                result = number1 / number2
            else:
                result = "Неизвестная операция"

            self.result_text.setText(str(result))

        except ValueError:
            self.result_text.setText("Введите корректные числа.")

    def insert_sum_to_cell(self):
        try:
            row_number = int(self.row_number_input.text())
            if 1 <= row_number <= 1048576:  # ограничение на номер строки в документе DOCX
                sum_result = float(self.result_text.text())  # предполагается, что result_text содержит численный результат
                cell_address = f"{row_number}"

                # Вставляем сумму в новую строку документа DOCX
                doc = Document()
                doc.add_paragraph(f"{cell_address}: {sum_result}")

                # Открываем диалоговое окно для ввода имени файла
                file_name, _ = QFileDialog.getSaveFileName(self, "Сохранить документ DOCX", "", "Word Files (*.docx);;All Files (*)")

                if file_name:
                    # Сохраняем документ DOCX
                    doc.save(file_name)
                    print(f"Вставлено в ячейку {cell_address}: {sum_result}. Результат сохранен в файл: {file_name}")
            else:
                print("Недопустимый номер строки.")
        except ValueError:
            print("Введите корректный номер строки.")

    def save_result_to_file(self):
        try:
            options = QFileDialog.Options()
            options |= QFileDialog.DontUseNativeDialog
            file_name, _ = QFileDialog.getSaveFileName(self, "Сохранить результат в файл", "", "Word Files (*.docx);;All Files (*)", options=options)

            if file_name:
                doc = Document()
                doc.add_paragraph(str(self.result_value))
                doc.save(file_name)
                print(f"Результат сохранен в файл DOCX: {file_name}")
        except Exception as e:
            print(f"Ошибка при сохранении в файл DOCX: {e}")

    def clear_data(self):
        self.result_text.clear()
        self.number1_input.clear()
        self.number2_input.clear()


#РАБОТА С ЧЕКБОКСАМИ
class AboutApp(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)

        self.init_ui()

    def init_ui(self):
        self.setWindowTitle('Справка')
        self.setFixedSize(520, 400)

        
        label = QLabel('Здесь вы можете получить всю необходимую информацию о приложении.<br>'
        'Узнайте текущую версию приложения, дату последнего обновления и ключевые<br> '
        'технические особенности. Мы постоянно работаем над улучшением функционала, <br>'
        'и здесь вы сможете найти подробности о последних изменениях.<br>'

        'Также в этом разделе вы найдете информацию о команде разработчиков,<br> '
        'ответственных за создание приложения. Узнайте больше о нашем опыте,<br> '
        'миссии и целях, чтобы быть в курсе всех тонкостей технологического<br> '
        'творчества.', self)
        layout = QVBoxLayout(self)
        layout.addWidget(label)


class AboutInstruction(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)

        self.init_ui()

    def init_ui(self):
        self.setWindowTitle('Инструкция')
        self.setFixedSize(500, 400)

        label = QLabel('анализ - калькулятор ипотечной ставки и кальуклятор прибыли', self)

        layout = QVBoxLayout(self)
        layout.addWidget(label)


#ЗАПУСК
if __name__ == "__main__":
    app = QApplication([])
    main_window = MainWindow()
    main_window.show()
    app.exec_()
